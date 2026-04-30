"""
규정집 DOCX/PDF → 청크 → ChromaDB 적재.

지원 문서:
  - mgmt: 4단계 BK21 관리운영지침 (제N조 형식, DOCX)
  - budget: 4단계 BK21 예산편성 및 집행기준 (제N조 형식, DOCX)
  - internal: 사업단 자체규정 (로마숫자/번호/괄호 계층, DOCX)
  - manual: 4단계 BK21 실무자 매뉴얼 ([대괄호섹션]/Ⅰ./1. 계층, PDF)

각 청크는 기존 bk21_qna 컬렉션에 doc_type=regulation 메타와 함께 upsert.
"""

import os
import re
import sys
import io
from dataclasses import dataclass, field
from typing import Iterator, Optional

import chromadb
from chromadb.utils import embedding_functions
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from dotenv import load_dotenv

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

load_dotenv()

REGULATIONS_DIR = "data/regulations"
CHROMA_DB_DIR = os.getenv("CHROMA_DB_DIR", "./data/chroma_db")
COLLECTION_NAME = "bk21_qna"
EMBEDDING_MODEL = "text-embedding-3-large"
BATCH_SIZE = 100

MIN_CHUNK_LEN = 40
MAX_CHUNK_LEN = 2000  # 초과 시 항(①②) 또는 호(1.) 단위로 분할

# 파일명 prefix / 키워드 → 문서 메타
# DOC_REGISTRY는 prefix(엄격 매칭) + 키워드(부분 매칭) 두 가지 키 지원
DOC_REGISTRY = {
    "mgmt":     {"name": "BK21 4단계 관리운영지침",      "short": "관리운영지침",      "style": "article"},
    "budget":   {"name": "BK21 4단계 예산편성 및 집행기준", "short": "예산편성·집행기준", "style": "article"},
    "internal": {"name": "사업단 자체규정",               "short": "사업단 자체규정",   "style": "outline"},
    "manual":   {"name": "BK21 4단계 실무자 매뉴얼",      "short": "실무자매뉴얼",      "style": "manual"},
}
# 파일명에 다음 키워드가 포함되면 해당 doc_key로 매핑 (prefix가 없는 파일용)
KEYWORD_TO_KEY = [
    ("매뉴얼", "manual"),
    ("실무자", "manual"),
]

# 조문 마커
RE_CHAPTER  = re.compile(r'^제\s*(\d+)\s*장\s*(.*)$')
RE_ARTICLE  = re.compile(r'^제\s*(\d+)\s*조(?:의\s*(\d+))?\s*\(([^)]+)\)\s*(.*)$')
RE_PARAGRAPH = re.compile(r'^([①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑰⑱⑲⑳])')

# 자체규정 마커
RE_ROMAN    = re.compile(r'^([ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ])\.\s*(.*)$')
RE_SECTION  = re.compile(r'^(\d+)\.\s+(.+?)(?:\t.*)?$')   # "1. 운영위원회"
RE_SUBSEC   = re.compile(r'^\((\d+)\)\s*(.+?)(?:\t.*)?$')  # "(1) 연구업적 성과급"

# 매뉴얼 마커
# 큰 섹션은 "[ 4단계 두뇌한국 (BK)21 사업 개요 ]" 같은 형식. UI 버튼([저장], [등록])은 너무 짧거나 동사형.
RE_MANUAL_TOP = re.compile(r'^\s*\[\s*([^\]]{6,80})\s*\]\s*$')
# "1) 가" 형식 (2차 소단원)
RE_NUMBERED_PAREN = re.compile(r'^\s*(\d+)\)\s+(.+?)$')

# 별표 — 두 가지 형식: "별표 1", "【별표 1】", "[별표 1]"
RE_APPENDIX = re.compile(r'^[【\[]?\s*별\s*표\s*(\d*)\s*[】\]]?\s*[\.:]?\s*(.*)$')

# 부칙 — 본문 제N조와 부칙 제N조의 ID 충돌 방지 위해 prefix 부여
RE_ADDENDUM = re.compile(r'^부\s*칙\s*(?:[<\(]\s*([\d\.\s]+)\s*[>\)])?')


@dataclass
class Chunk:
    doc_key: str           # mgmt / budget / internal
    chapter: str = ""      # 제1장 / Ⅰ.
    chapter_title: str = ""
    article_no: str = ""   # 제1조 / 1. / 별표1
    article_title: str = ""
    sub_no: str = ""       # 제1조의2 / (1)
    sub_title: str = ""
    paragraphs: list = field(default_factory=list)  # 본문 단락 또는 표(md)
    chunk_id_suffix: str = ""

    def text(self) -> str:
        body = "\n".join(self.paragraphs).strip()
        return body

    def header_lines(self, doc_meta: dict) -> str:
        """LLM이 이 청크의 출처를 알 수 있도록 헤더 텍스트 구성."""
        lines = [doc_meta["name"]]
        if self.chapter:
            tail = f" {self.chapter_title}" if self.chapter_title else ""
            lines.append(f"{self.chapter}{tail}")
        if self.article_no:
            tail = f" ({self.article_title})" if self.article_title else ""
            lines.append(f"{self.article_no}{tail}")
        if self.sub_no:
            tail = f" {self.sub_title}" if self.sub_title else ""
            lines.append(f"  {self.sub_no}{tail}")
        return "\n".join(lines)

    def citation(self, doc_meta: dict) -> str:
        """답변에서 인용할 짧은 형식.
        - article style: "관리운영지침 제12조 (참여대학원생 자격)"
        - outline style: "사업단 자체규정 Ⅱ.1. 운영위원회"
        - manual style: "실무자매뉴얼 [교육연구단(팀) 관리 및 운영] Ⅱ.3. 참여대학원생"
        """
        parts = [doc_meta["short"]]
        if doc_meta["style"] == "manual":
            if self.chapter_title:
                parts.append(f"[{self.chapter_title}]")
            ref = ""
            if self.article_no:
                ref += self.article_no
            if self.sub_no:
                ref += self.sub_no
            if ref:
                parts.append(ref)
            title = self.sub_title or self.article_title
            if title:
                parts.append(title[:35])
        elif doc_meta["style"] == "outline":
            ref = ""
            if self.chapter:
                ref += self.chapter
            if self.article_no:
                ref += self.article_no
            if self.sub_no:
                ref += self.sub_no
            if ref:
                parts.append(ref)
            title = self.sub_title or self.article_title or self.chapter_title
            if title:
                parts.append(title[:30])
        else:
            if self.article_no:
                parts.append(self.article_no)
            if self.sub_no:
                parts.append(self.sub_no)
            if self.article_title and not self.article_no.startswith('별표'):
                parts.append(f"({self.article_title})")
        return " ".join(parts)

    def chroma_id(self, doc_key: str) -> str:
        # ID 충돌 방지: reg_<doc>_<chapter>_<art>_<sub>_<title>_<suffix>
        slug = lambda s: re.sub(r'[^\w가-힣]+', '', s)
        bits = [doc_key]
        if self.chapter:
            bits.append(slug(self.chapter))
        if self.article_no:
            bits.append(slug(self.article_no))
        if self.sub_no:
            bits.append(slug(self.sub_no))
        # outline style은 같은 번호가 여러번 나타날 수 있으므로 title slug 포함
        title = self.sub_title or self.article_title
        if title:
            bits.append(slug(title)[:20])
        if self.chunk_id_suffix:
            bits.append(self.chunk_id_suffix)
        return "reg_" + "_".join(bits)


def iter_body(doc: Document) -> Iterator[tuple]:
    """문서 본문을 순서대로 (kind, obj) 튜플로 산출. kind = 'p' | 'tbl'."""
    body = doc.element.body
    para_idx = 0
    table_idx = 0
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            yield ('p', doc.paragraphs[para_idx])
            para_idx += 1
        elif child.tag == qn('w:tbl'):
            yield ('tbl', doc.tables[table_idx])
            table_idx += 1


def table_to_markdown(table: Table) -> str:
    """표를 markdown으로 변환. 셀 내 \\n은 공백으로, | 는 escape."""
    rows = []
    seen = set()
    for ri, row in enumerate(table.rows):
        cells_raw = []
        for c in row.cells:
            cid = id(c._tc)
            if cid in seen:
                cells_raw.append("")  # merged cell continuation
            else:
                seen.add(cid)
                txt = c.text.replace('|', '\\|').replace('\n', ' ').strip()
                cells_raw.append(txt)
        rows.append("| " + " | ".join(cells_raw) + " |")
        if ri == 0:
            rows.append("|" + "|".join(["---"] * len(cells_raw)) + "|")
    return "\n".join(rows)


def is_toc_or_metadata(text: str) -> bool:
    """목차/개정 이력/표지 라인 판별 — 본문 시작 전까지 스킵용."""
    t = text.strip()
    if not t:
        return True
    # 목차: 탭 뒤 페이지 번호("1", "1-2", "3-7" 등)
    if '\t' in t and re.search(r'\t[\d\-]+\s*$', t):
        return True
    # 개정/제정 이력
    if re.match(r'^\d{4}\.\s*\d{1,2}\.', t) and ('개정' in t or '제정' in t):
        return True
    if re.match(r'^(개정|제정):\s*\d{4}', t):
        return True
    # 표지
    if t in ('- 목 차 -', '목 차', '목차'):
        return True
    return False


def parse_article_style(doc: Document, doc_key: str, version: str) -> list:
    """mgmt/budget 형식 파서 — 제N조 단위 청킹."""
    chunks: list = []
    cur = Chunk(doc_key=doc_key)
    body_started = False
    addendum_idx = 0   # 부칙 회수 — 1, 2, 3...
    addendum_prefix = ""  # 부칙(2025.12.24) 형식

    def flush():
        if cur.text().strip() or cur.paragraphs:
            chunks.append(cur)

    for kind, obj in iter_body(doc):
        if kind == 'p':
            text = obj.text.strip()
            if not body_started:
                # 첫 "제 N 장" 이전은 모두 스킵
                if RE_CHAPTER.match(re.sub(r'\s+', '', text)) or RE_CHAPTER.match(text):
                    body_started = True
                else:
                    continue
            if not text:
                continue

            # 부칙 시작 — 후속 제N조의 prefix 갱신
            m = RE_ADDENDUM.match(text)
            if m:
                flush()
                addendum_idx += 1
                date = m.group(1)
                addendum_prefix = f"부칙({date.strip()})" if date else f"부칙{addendum_idx}"
                cur = Chunk(
                    doc_key=doc_key,
                    chapter=addendum_prefix,
                    chapter_title="",
                )
                continue

            # 별표 — 새 청크 (【별표 N】 / 별표 N 모두)
            m = RE_APPENDIX.match(text)
            if m and (text.startswith('별') or text.startswith('【') or text.startswith('[')):
                flush()
                num = m.group(1) or ''
                cur = Chunk(
                    doc_key=doc_key,
                    article_no=f"별표{num}",
                    article_title=m.group(2).strip(),
                )
                continue

            m = RE_CHAPTER.match(text)
            if m:
                # 새 장 시작
                flush()
                cur = Chunk(
                    doc_key=doc_key,
                    chapter=f"제{m.group(1)}장",
                    chapter_title=m.group(2).strip(),
                )
                continue

            m = RE_ARTICLE.match(text)
            if m:
                # 새 조 시작 — 이전 조 flush, 새 조의 헤더만 가져와 사용
                prev_chapter = cur.chapter
                prev_chapter_title = cur.chapter_title
                flush()
                article_no = f"제{m.group(1)}조"
                if m.group(2):
                    article_no += f"의{m.group(2)}"
                # 부칙 안의 조는 prefix 부여
                if addendum_prefix:
                    article_no = f"{addendum_prefix} {article_no}"
                cur = Chunk(
                    doc_key=doc_key,
                    chapter=prev_chapter,
                    chapter_title=prev_chapter_title,
                    article_no=article_no,
                    article_title=m.group(3).strip(),
                )
                tail = m.group(4).strip()
                if tail:
                    cur.paragraphs.append(tail)
                continue

            # 일반 본문 (항/호/평문) — 현재 청크에 추가
            cur.paragraphs.append(text)

        elif kind == 'tbl':
            if not body_started:
                continue
            md = table_to_markdown(obj)
            if md.strip():
                cur.paragraphs.append(md)

    flush()
    return chunks


def parse_outline_style(doc: Document, doc_key: str, version: str) -> list:
    """internal 형식 파서 — 로마.번호.(괄호) 계층."""
    chunks: list = []
    cur = Chunk(doc_key=doc_key)
    body_started = False
    seen_first_roman = False  # 첫 번째 Ⅰ.은 목차, 두 번째가 본문

    def flush():
        if cur.text().strip() or cur.paragraphs:
            chunks.append(cur)

    cur_chapter = ""
    cur_chapter_title = ""
    cur_section = ""
    cur_section_title = ""

    for kind, obj in iter_body(doc):
        if kind == 'p':
            text = obj.text.strip()
            if is_toc_or_metadata(text):
                continue
            if not text:
                continue

            # 로마숫자 (Ⅰ. ...)
            m = RE_ROMAN.match(text)
            if m:
                if not seen_first_roman:
                    seen_first_roman = True
                    # 목차의 첫 Ⅰ.인지 본문의 첫 Ⅰ.인지 모르므로 일단 본문으로 가정
                if not body_started:
                    body_started = True
                flush()
                cur_chapter = m.group(1) + "."
                cur_chapter_title = m.group(2).strip()
                cur_section = ""
                cur_section_title = ""
                cur = Chunk(
                    doc_key=doc_key,
                    chapter=cur_chapter,
                    chapter_title=cur_chapter_title,
                )
                continue

            if not body_started:
                continue

            # 1. 운영위원회 (탭 + 페이지 번호 없는 본문)
            m = RE_SECTION.match(text)
            if m and not (m.group(2).rstrip().endswith(tuple('0123456789-'))):
                flush()
                cur_section = f"{m.group(1)}."
                cur_section_title = m.group(2).strip()
                cur = Chunk(
                    doc_key=doc_key,
                    chapter=cur_chapter,
                    chapter_title=cur_chapter_title,
                    article_no=cur_section,
                    article_title=cur_section_title,
                )
                continue

            # (1) 연구업적 성과급
            m = RE_SUBSEC.match(text)
            if m:
                flush()
                cur = Chunk(
                    doc_key=doc_key,
                    chapter=cur_chapter,
                    chapter_title=cur_chapter_title,
                    article_no=cur_section,
                    article_title=cur_section_title,
                    sub_no=f"({m.group(1)})",
                    sub_title=m.group(2).strip(),
                )
                continue

            # 평문 — 현재 청크에 추가
            cur.paragraphs.append(text)

        elif kind == 'tbl':
            if not body_started:
                continue
            md = table_to_markdown(obj)
            if md.strip():
                cur.paragraphs.append(md)

    flush()
    return chunks


def _split_by_marker(lines: list, marker_re) -> list:
    """주어진 정규식 마커 라인 시작점으로 그룹 분할."""
    groups, buf = [], []
    for ln in lines:
        if marker_re.match(ln) and buf:
            groups.append(buf)
            buf = [ln]
        else:
            buf.append(ln)
    if buf:
        groups.append(buf)
    return groups


def _split_by_length(lines: list, max_len: int) -> list:
    """길이 기준으로 청크 묶음을 더 잘게 (단락 경계 보존)."""
    groups, buf, cur_len = [], [], 0
    for ln in lines:
        if cur_len + len(ln) > max_len and buf:
            groups.append(buf)
            buf = [ln]
            cur_len = len(ln)
        else:
            buf.append(ln)
            cur_len += len(ln)
    if buf:
        groups.append(buf)
    return groups


# 매뉴얼 본문 분할용 마커: ❚ 와 ➊~➓
RE_MANUAL_HDR = re.compile(r'^\s*❚')
RE_MANUAL_NUM = re.compile(r'^\s*[➊➋➌➍➎➏➐➑➒➓]')


def split_oversized(chunks: list) -> list:
    """청크가 MAX_CHUNK_LEN 초과면 단계적으로 분할:
       1. ① 항 마커 → 2. ❚ 마커 → 3. ➊ 마커 → 4. 길이 기준 강제 분할
    """
    out = []
    for ch in chunks:
        body = ch.text()
        if len(body) <= MAX_CHUNK_LEN:
            out.append(ch)
            continue

        lines = body.split('\n')
        groups = None
        for marker in (RE_PARAGRAPH, RE_MANUAL_HDR, RE_MANUAL_NUM):
            cand = _split_by_marker(lines, marker)
            if len(cand) > 1:
                groups = cand
                break

        if groups is None:
            # 어떤 마커로도 분할 안 되면 길이 기준
            groups = _split_by_length(lines, MAX_CHUNK_LEN)

        # 분할 결과가 여전히 너무 크면 길이 기준 추가 분할
        final_groups = []
        for g in groups:
            g_text_len = sum(len(l) for l in g)
            if g_text_len > MAX_CHUNK_LEN:
                final_groups.extend(_split_by_length(g, MAX_CHUNK_LEN))
            else:
                final_groups.append(g)

        for i, g in enumerate(final_groups):
            new = Chunk(**{
                k: v for k, v in ch.__dict__.items()
                if k not in ('paragraphs', 'chunk_id_suffix')
            })
            new.paragraphs = g
            new.chunk_id_suffix = f"p{i+1}"
            # 항 번호를 sub_no로 보강
            first = g[0] if g else ""
            mp = RE_PARAGRAPH.match(first)
            if mp and not new.sub_no:
                new.sub_no = mp.group(1) + "항"
            out.append(new)
    return out


def filter_short(chunks: list) -> list:
    """본문이 너무 짧은 청크 제거."""
    return [c for c in chunks if len(c.text()) >= MIN_CHUNK_LEN]


def parse_doc_filename(fname: str) -> tuple:
    """파일명 → (doc_key, version).
    - regulations/<key>_<displayname>_<YYYY.MM>.<ext>: prefix 직접 사용
    - 그 외: KEYWORD_TO_KEY 휴리스틱
    """
    base = os.path.basename(fname).rsplit(".", 1)[0]
    parts = base.split("_")
    key = parts[0] if parts and parts[0] in DOC_REGISTRY else None

    if key is None:
        for keyword, mapped in KEYWORD_TO_KEY:
            if keyword in base:
                key = mapped
                break

    version_match = re.search(r'(\d{4}\.\d{1,2})', base)
    version = version_match.group(1) if version_match else ""
    return key, version


def read_pdf_pages(path: str) -> list:
    """PyMuPDF + sort=True로 페이지별 텍스트 산출. 매뉴얼 같은 PDF용."""
    import fitz
    doc = fitz.open(path)
    pages = [doc[i].get_text(sort=True) for i in range(len(doc))]
    doc.close()
    return pages


def _normalize_chapter_label(label: str) -> str:
    """PDF 추출에서 디지트 분리로 생긴 '4단계 두뇌한국 (BK)21' 류의 빈 괄호/공백 정리."""
    s = re.sub(r'\(\s+\)', '', label)  # 빈 괄호
    s = re.sub(r'\s{2,}', ' ', s).strip()
    return s


def parse_manual_style(pages: list, doc_key: str, version: str) -> list:
    """매뉴얼 PDF 파서.

    PDF 추출 과정에서 디지트가 자주 분리되어 (예: "3. 사업비..." → "3 사업비...")
    Ⅰ./1. 마커 검출이 신뢰 불가. [큰섹션] 4개 단위로만 분리하고, 그 안은 split_oversized
    에서 ❚ ➊ 마커 + 길이 기준으로 잘게 분할.
    """
    chunks: list = []
    cur_chapter = ""
    cur_chapter_title = ""
    cur_lines: list = []
    body_started = False

    def flush():
        if not cur_lines:
            return
        body = "\n".join(cur_lines).strip()
        if not body:
            return
        chunks.append(Chunk(
            doc_key=doc_key,
            chapter=cur_chapter,
            chapter_title=cur_chapter_title,
            paragraphs=list(cur_lines),
        ))

    SKIP_LABELS = (
        "내정보", "참고", "저장", "등록", "확인", "다운로드", "발급",
        "다음", "돌아가기", "검색", "클릭", "보기", "선택", "닫기",
    )

    all_text = "\n".join(pages)
    for raw_line in all_text.split("\n"):
        text = raw_line.strip()
        if not text:
            continue
        if text.startswith("|") and text.endswith("|"):
            continue

        m = RE_MANUAL_TOP.match(text)
        if m:
            label = _normalize_chapter_label(m.group(1).strip())
            if len(label) < 6:
                continue
            if any(skip in label for skip in SKIP_LABELS):
                continue
            if label != cur_chapter:
                flush()
                cur_lines = []
                cur_chapter = label
                cur_chapter_title = label
                body_started = True
            continue

        if not body_started:
            continue
        cur_lines.append(text)

    flush()
    return chunks


def build_chunks_for_file(path: str) -> list:
    doc_key, version = parse_doc_filename(path)
    if not doc_key or doc_key not in DOC_REGISTRY:
        print(f"[!] 알 수 없는 문서 prefix/키워드: {path}")
        return []
    meta = DOC_REGISTRY[doc_key]

    if meta["style"] == "manual":
        if not path.lower().endswith(".pdf"):
            print(f"[!] manual 스타일은 PDF만 지원: {path}")
            return []
        pages = read_pdf_pages(path)
        chunks = parse_manual_style(pages, doc_key, version)
    else:
        if not path.lower().endswith(".docx"):
            print(f"[!] {meta['style']} 스타일은 DOCX만 지원: {path}")
            return []
        doc = Document(path)
        if meta["style"] == "article":
            chunks = parse_article_style(doc, doc_key, version)
        else:
            chunks = parse_outline_style(doc, doc_key, version)

    chunks = split_oversized(chunks)
    chunks = filter_short(chunks)
    for c in chunks:
        c._source_path = path
        c._version = version
    return chunks


def upsert_chunks(chunks: list, source_paths: list):
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("[!] OPENAI_API_KEY 미설정")
        return

    client = chromadb.PersistentClient(path=CHROMA_DB_DIR)
    openai_ef = embedding_functions.OpenAIEmbeddingFunction(
        api_key=api_key,
        model_name=EMBEDDING_MODEL,
    )
    collection = client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=openai_ef,
    )

    # 기존 reg_* ID 모두 제거 (개정 시 stale 데이터 방지)
    existing = collection.get(where={"doc_type": "regulation"})
    if existing and existing.get("ids"):
        print(f"[-] 기존 규정 청크 {len(existing['ids'])}건 제거")
        collection.delete(ids=existing["ids"])

    documents, metadatas, ids = [], [], []
    for ch in chunks:
        meta = DOC_REGISTRY[ch.doc_key]
        header = ch.header_lines(meta)
        body = ch.text()
        embed_text = f"{header}\n\n{body}"
        chroma_meta = {
            "doc_type": "regulation",
            "doc_key": ch.doc_key,
            "doc_name": meta["name"],
            "version": ch._version,
            "chapter": f"{ch.chapter} {ch.chapter_title}".strip(),
            "article_no": ch.article_no,
            "article_title": ch.article_title,
            "sub_no": ch.sub_no,
            "sub_title": ch.sub_title,
            "citation": ch.citation(meta),
            "source_path": ch._source_path,
        }
        documents.append(embed_text)
        metadatas.append(chroma_meta)
        ids.append(ch.chroma_id(ch.doc_key))

    # 중복 ID 검출 (안전장치)
    if len(set(ids)) != len(ids):
        from collections import Counter
        dups = [i for i, c in Counter(ids).items() if c > 1]
        print(f"[!] 중복 ID {len(dups)}건. 첫 5개: {dups[:5]}")
        # 중복은 suffix로 회피
        seen = {}
        for i, _id in enumerate(ids):
            seen[_id] = seen.get(_id, 0) + 1
            if seen[_id] > 1:
                ids[i] = f"{_id}_dup{seen[_id]}"

    print(f"[-] 임베딩+upsert 시작: {len(documents)}건 (배치 {BATCH_SIZE})")
    for i in range(0, len(documents), BATCH_SIZE):
        collection.upsert(
            documents=documents[i:i+BATCH_SIZE],
            metadatas=metadatas[i:i+BATCH_SIZE],
            ids=ids[i:i+BATCH_SIZE],
        )
        print(f"    {min(i+BATCH_SIZE, len(documents))}/{len(documents)}")

    print(f"[v] 컬렉션 총 문서: {collection.count()}건")


def main():
    paths = sorted(
        os.path.join(REGULATIONS_DIR, f)
        for f in os.listdir(REGULATIONS_DIR)
        if (f.endswith(".docx") or f.endswith(".pdf")) and not f.startswith("~")
    )
    if not paths:
        print(f"[!] {REGULATIONS_DIR}/ 에 DOCX 없음")
        return

    all_chunks = []
    for p in paths:
        print(f"\n=== 파싱: {p} ===")
        chunks = build_chunks_for_file(p)
        print(f"[-] {len(chunks)}개 청크 생성")
        # 통계
        lens = [len(c.text()) for c in chunks]
        if lens:
            print(f"    길이 평균/최소/최대/p25/p75: {sum(lens)//len(lens)} / {min(lens)} / {max(lens)} / {sorted(lens)[len(lens)//4]} / {sorted(lens)[3*len(lens)//4]}")
        # 첫 3개 미리보기
        for c in chunks[:3]:
            meta = DOC_REGISTRY[c.doc_key]
            print(f"    --- {c.citation(meta)} ---")
            print(f"      {c.text()[:150]}")
        all_chunks.extend(chunks)

    print(f"\n=== 총 {len(all_chunks)}개 청크 적재 시작 ===")
    upsert_chunks(all_chunks, paths)


if __name__ == "__main__":
    main()
